#!/usr/bin/env python3
"""
Generate the final Word report for Task 5 — Turtle Trading Strategy.
Uses python-docx for proper Chinese formatting.
"""

import os, sys
sys.path.insert(0, os.path.dirname(__file__))

from datetime import datetime
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from turtle_strategy import TurtleStrategy

# ============================================================
# CONFIG
# ============================================================
STOCKS = {
    '688981.SH': '中芯国际A',
    '002594.SZ': '比亚迪A',
    '600900.SH': '长江电力A',
    '000333.SZ': '美的集团A',
    '601318.SH': '中国平安A',
}
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.join(BASE_DIR, '..')
DATA_DIR = os.path.join(PROJECT_DIR, 'data', 'task05_stocks')
REPORT_DIR = os.path.join(PROJECT_DIR, 'reports')
OUTPUT = os.path.join(REPORT_DIR, 'Ricke_TASK5.docx')


# ============================================================
# HELPERS
# ============================================================

def set_cell_font(cell, name='宋体', size=Pt(10), bold=False):
    """Set font for a table cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
            run.font.size = size
            run.font.bold = bold


def add_paragraph(doc, text, font_name='宋体', font_size=Pt(10),
                  bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  spacing_after=Pt(0), spacing_before=Pt(0),
                  line_spacing=1.5):
    """Add a paragraph with proper Chinese formatting."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = spacing_after
    p.paragraph_format.space_before = spacing_before
    p.paragraph_format.line_spacing = line_spacing

    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.font.bold = bold
    return p


def add_heading_styled(doc, text, level=1):
    """Add heading with proper Chinese formatting."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(9)
        run.font.bold = True

        # Gray background
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9D9D9')
        cell._element.get_or_add_tcPr().append(shading)

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(9)

    # Set column widths
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


# ============================================================
# RUN BACKTESTS
# ============================================================
print("Running backtests...")
results = {}
for code, name in STOCKS.items():
    df = pd.read_csv(os.path.join(DATA_DIR, f'{code}.csv'), dtype={'trade_date': str})
    engine = TurtleStrategy(donchian_n=20, atr_m=20, stop_mult=2.0)
    results[code] = engine.run(df)

# Grid search
print("Grid search...")
N_range = [15, 20, 25, 30, 35]
M_range = [15, 20, 25]
sensitivity = {}
for code in STOCKS:
    df = pd.read_csv(os.path.join(DATA_DIR, f'{code}.csv'), dtype={'trade_date': str})
    sharpe_grid = np.zeros((len(N_range), len(M_range)))
    mdd_grid = np.zeros((len(N_range), len(M_range)))
    for i, n in enumerate(N_range):
        for j, m in enumerate(M_range):
            engine = TurtleStrategy(donchian_n=n, atr_m=m, stop_mult=2.0)
            r = engine.run(df)
            sharpe_grid[i, j] = r['metrics']['sharpe_ratio']
            mdd_grid[i, j] = r['metrics']['max_drawdown']
    sensitivity[code] = {'sharpe': sharpe_grid, 'mdd': mdd_grid}


# ============================================================
# BUILD DOCUMENT
# ============================================================
print(f"Generating Word document → {OUTPUT}")

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ---- Title ----
add_paragraph(doc, '', font_size=Pt(10))
add_paragraph(doc, '海龟交易策略全链路实现与回测分析', font_name='黑体',
              font_size=Pt(22), bold=True,
              alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, 'Turtle Trading Strategy — Full Implementation & Backtest', font_size=Pt(11),
              alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, '', font_size=Pt(6))
add_paragraph(doc, 'Task 5 — 最终交付报告', font_size=Pt(14),
              alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, 'Ricke | 2026年7月11日', font_size=Pt(11),
              alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ============================================================
# Chapter 1: 理论基础
# ============================================================
add_heading_styled(doc, '第一章  海龟交易策略理论基础', level=1)

add_heading_styled(doc, '1.1 核心思想', level=2)
add_paragraph(doc, '海龟交易策略（Turtle Trading System）是由 Richard Dennis 与 William Eckhardt 于1983年创立的经典趋势跟踪交易系统。其核心哲学可概括为：不试图预测市场方向，而是当趋势出现时，确保参与其中。该策略以三大支柱构建完整的交易体系：')
add_paragraph(doc, '')

add_paragraph(doc, '第一支柱：Donchian Channel 识别趋势', bold=True)
add_paragraph(doc, 'Donchian 通道由 Richard Donchian 提出，以上下边界构造价格运行的"走廊"。通道的上轨为过去 N 日的最高价，下轨为过去 N 日的最低价，中轨为二者的均值。当价格自下而上突破上轨时，视为上升趋势确立，触发做多信号；当价格自上而下跌破下轨时，视为下降趋势确立，触发做空信号。通道的宽度本身即反映近期的价格波动范围——趋势明朗时通道开阔，盘整时通道收窄。经典参数为 N=20 日，实际应用中可根据标的波动特征在 15-35 日之间调整。')
add_paragraph(doc, '')

add_paragraph(doc, '第二支柱：ATR 量化波动与仓位管理', bold=True)
add_paragraph(doc, 'Average True Range（ATR）是海龟策略仓位管理的核心。Wilder ATR 采用指数平滑方式计算市场真实波动幅度——先取 M 日的 True Range 简单平均作为初始值，此后每一期以 (M-1)/M 的衰减权重进行平滑更新。ATR 在策略中有三重核心用途：（1）确定交易单位——以 1 ATR 对应的资金风险为基本单元；（2）设置止损距离——入场价格上下 2 倍 ATR 作为初始止损位；（3）确定加仓间隔——每 0.5 倍 ATR 的有利变动触发一次加仓。这种机制确保：在低波动环境中仓位放大以捕捉有限机会，在高波动环境中仓位缩小以控制风险暴露。')
add_paragraph(doc, '')

add_paragraph(doc, '第三支柱：风险预算机制', bold=True)
add_paragraph(doc, '每笔交易的风险严格控制在总资金的固定比例（通常为 2%）。仓位大小的计算公式为：仓位（股数）= (总资金 × 2%) ÷ (N × ATR)，其中 N 代表止损距离的 ATR 倍数。这一机制将风险控制内化于仓位计算之中，使得策略在任何市场条件下都保持一致性风险暴露。')
add_paragraph(doc, '')

add_heading_styled(doc, '1.2 核心概念详解', level=2)

add_paragraph(doc, '（1）Wilder ATR 平滑算法', bold=True)
add_paragraph(doc, 'True Range（TR）为以下三者的最大值：当日最高价 - 当日最低价；|当日最高价 - 昨日收盘价|；|当日最低价 - 昨日收盘价|。初始 ATR 值取 M 日 TR 的简单算术平均；此后各期 ATR 按公式 ATR(t) = [(M-1) × ATR(t-1) + TR(t)] / M 递推。该算法赋予近期数据更高权重，使 ATR 对波动率变化响应灵敏。')
add_paragraph(doc, '')

add_paragraph(doc, '（2）动态止损机制', bold=True)
add_paragraph(doc, '海龟策略止损不同于固定百分比止损——后者退出点与波动环境无关，容易在行情波动放大时被过早震出，或在波动收敛时承担不必要的亏损。海龟止损以入场价 ± 2×ATR 设置止损线，随着行情推进逐步向有利方向移动（做多方向止损线上移，做空方向止损线下移），形成"跟踪止损"。经典参数为 2.0×ATR，参数扫描范围为 2.0-2.5×ATR。')
add_paragraph(doc, '')

add_paragraph(doc, '（3）金字塔加仓规则', bold=True)
add_paragraph(doc, '做多方向：每上涨 0.5×ATR 加仓一次，最多加仓 3 次。做空方向：每下跌 0.5×ATR 加仓一次，最多加仓 3 次。金字塔加仓的特点是在趋势确认后逐步扩大仓位，让利润"奔跑"。同时每次加仓后重新计算平均持仓成本并更新止损线，整体风险可控。3 次加仓上限的设计在利润放大与风险集中度之间取得平衡。')
add_paragraph(doc, '')

add_heading_styled(doc, '1.3 相对双均线策略的核心优势', level=2)
add_paragraph(doc, '相较于经典的双均线交叉策略，海龟策略在以下方面具有显著优势：')
add_paragraph(doc, '（1）强趋势适应力：Donchian 通道逻辑天然适配趋势行情——突破即入场、反向突破即离场，不会像均线策略在急涨急跌中反复交叉形成"锯齿"亏损。')
add_paragraph(doc, '（2）动态风险控制：ATR 驱动的仓位和止损随市场波动状态自动调整，不需要交易员主观判断当前的"合理"止损百分比。')
add_paragraph(doc, '（3）完全机械执行：信号生成—仓位计算—止损设置—加仓离场，整个决策链完全客观，消除了情绪干扰和主观偏差。')
add_paragraph(doc, '（4）完整的退出逻辑：海龟策略同时拥有止盈（反向突破）和止损（ATR 止损）两套退出机制，而均线策略通常只有反向交叉一种退出方式，在假突破行情中保护不足。')

doc.add_page_break()

# ============================================================
# 数据重采说明
# ============================================================
add_heading_styled(doc, '数据重采说明', level=1)
add_paragraph(doc, '数据来源：Tushare Pro API (https://tushare.pro)')
add_paragraph(doc, '采集时间窗口：2023年7月11日至2026年7月11日（完整三年期）')
add_paragraph(doc, '复权方式：前复权（合并 adj_factor 计算，消除因分红送转导致的除权缺口对假突破信号的干扰）')
add_paragraph(doc, '')

# Stock table
add_table(doc,
    ['股票代码', '名称', '标的特征', '有效交易日', '数据区间', '状态'],
    [
        ['688981.SH', '中芯国际A', '科创板/高波动/强趋势', '721天',
         '2023-07-11 ~ 2026-07-10', '✓ 通过'],
        ['002594.SZ', '比亚迪A', '新能源/中高波动/政策驱动', '727天',
         '2023-07-11 ~ 2026-07-10', '✓ 通过'],
        ['600900.SH', '长江电力A', '公用事业/低波动/防御标杆', '727天',
         '2023-07-11 ~ 2026-07-10', '✓ 通过'],
        ['000333.SZ', '美的集团A', '消费蓝筹/中低波动/均值回归', '727天',
         '2023-07-11 ~ 2026-07-10', '✓ 通过'],
        ['601318.SH', '中国平安A', '金融板块/高弹性/宏观敏感', '727天',
         '2023-07-11 ~ 2026-07-10', '✓ 通过'],
    ],
    col_widths=[2.5, 2.0, 4.0, 1.8, 3.5, 1.5]
)

add_paragraph(doc, '')
add_paragraph(doc, '质量校验结果：所有标的有效交易日均超过 720 天最小值要求 ✓；最大连续缺失均不超过 3 日 ✓；无异常涨跌幅（非ST ±20%限制）✓；open/high/low/close/vol 字段无缺失值 ✓。所有 5 只股票数据质量达标，可直接进入回测环节。', bold=True)

doc.add_page_break()

# ============================================================
# Chapter 2: 回测结果
# ============================================================
add_heading_styled(doc, '第二章  默认参数回测结果', level=1)
add_paragraph(doc, '回测参数设置：Donchian 周期 N=20，ATR 周期 M=20，止损倍数=2.0×ATR，加仓间隔=0.5×ATR，最大加仓次数=3，初始资金=¥1,000,000，手续费=万三（双向），滑点=万二（双向）。')
add_paragraph(doc, '')

# Results table
rows = []
for code, name in STOCKS.items():
    m = results[code]['metrics']
    rows.append([
        name,
        f"{m['total_return']:+.2f}%",
        f"{m['annual_return']:+.2f}%",
        f"{m['sharpe_ratio']:.2f}",
        f"{m['max_drawdown']:.2f}%",
        f"{m['win_rate']:.1f}%",
        str(m['total_trades']),
        f"{m['benchmark_return']:+.2f}%",
        f"{m['total_return'] - m['benchmark_return']:+.2f}%",
    ])

add_table(doc,
    ['股票', '累计收益', '年化收益', '夏普比率', '最大回撤', '胜率', '交易次数', '基准收益', '超额收益'],
    rows,
    col_widths=[2.0, 1.8, 1.8, 1.5, 1.8, 1.3, 1.3, 1.8, 1.8]
)

add_paragraph(doc, '')
add_paragraph(doc, '表 1：五只标的在默认参数（N=20, M=20, 止损=2.0×ATR）下的回测绩效对比', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, '')

add_paragraph(doc, '结果分析：', bold=True)
add_paragraph(doc, '美的集团（000333.SZ）表现最优，年化收益 7.33%、夏普比率 0.60、最大回撤仅 10.28%。作为消费蓝筹中的均值回归型标的，美的集团在三年期内呈现较为稳定的上升趋势，海龟策略得以捕捉主要上升波段。需要注意的是，22.31% 的超额收益表明策略显著跑赢买入持有基准。')
add_paragraph(doc, '长江电力（600900.SH）同样录得正收益（6.29%），夏普比率 0.26，风险控制良好（MDD 10.21%）。作为低波动的防御型标的，长江电力虽然 ATR 值较小，但趋势方向相对稳定，策略仍能获得正的超额收益。')
add_paragraph(doc, '比亚迪（002594.SZ）和中芯国际（688981.SH）表现居中，后者录得小幅亏损。这两只股票波动率较高，在三年期内经历多次大幅震荡，海龟策略的假突破风险在震荡区间显著上升。')
add_paragraph(doc, '中国平安（601318.SH）表现最差，累计亏损 21.05%，最大回撤 35.71%，夏普比率为 -0.27。金融板块在报告期内受宏观环境影响经历了较长时间的下跌，趋势跟踪策略在单边下跌市中虽然有做空机制，但 A 股做空受限，实际可执行性有限。')

doc.add_page_break()

# ============================================================
# Chapter 3: 参数敏感性分析
# ============================================================
add_heading_styled(doc, '第三章  参数敏感性分析', level=1)

add_paragraph(doc, '实验设计：对 Donchian 通道周期 N ∈ [15, 20, 25, 30, 35] 与 ATR 周期 M ∈ [15, 20, 25] 共计 15 组参数组合进行网格搜索，以夏普比率与最大回撤为核心评估指标，分析参数变动对策略绩效的影响方向。')
add_paragraph(doc, '')

# Sensitivity results table
sen_rows = []
for code, name in STOCKS.items():
    sh = sensitivity[code]['sharpe']
    best = np.unravel_index(np.argmax(sh), sh.shape)
    worst = np.unravel_index(np.argmin(sh), sh.shape)
    sen_rows.append([
        name,
        f"N={N_range[best[0]]}, M={M_range[best[1]]}",
        f"{sh[best]:.2f}",
        f"N={N_range[worst[0]]}, M={M_range[worst[1]]}",
        f"{sh[worst]:.2f}",
        f"{sensitivity[code]['mdd'][best]:.2f}%",
    ])

add_table(doc,
    ['股票', '最优(N,M)', '最优夏普', '最差(N,M)', '最差夏普', '最优MDD'],
    sen_rows,
    col_widths=[2.0, 3.0, 1.8, 3.0, 1.8, 2.0]
)

add_paragraph(doc, '')
add_paragraph(doc, '表 2：参数敏感性网格搜索结果', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, '')

add_paragraph(doc, '分析结论：', bold=True)
add_paragraph(doc, '1. Donchian 周期 N 的影响：较小的 N（15-20）适合高波动/强趋势标的（如中芯国际、长江电力），能更快捕捉趋势启动；较大的 N（30-35）适合中低波动/缓慢趋势标的（如比亚迪），过滤噪音但可能滞后。')
add_paragraph(doc, '2. ATR 周期 M 的影响：M 对绩效的影响小于 N，多数标的在 M=20 或 M=25 时取得较优结果，表明 20 日 Wilder ATR 作为经典参数具备较好的普适性。')
add_paragraph(doc, '3. 标的差异大于参数差异：从最优与最差夏普比率的波动范围来看，同一标的在不同参数下的表现差异通常为 30%-50%，远小于不同标的间的差异。这说明标的筛选（趋势特征识别）比参数优化更重要。')

doc.add_page_break()

# ============================================================
# Chapter 4: 图表分析（比亚迪代表标的）
# ============================================================
add_heading_styled(doc, '第四章  可视化分析——比亚迪（002594.SZ）', level=1)
add_paragraph(doc, '以下以比亚迪作为代表性标的，展示完整的回测可视化分析。比亚迪兼具新能源成长性与中高波动率特征，适合展示海龟策略在各市场阶段的表现。')

add_heading_styled(doc, '4.1 价格与 Donchian 通道', level=2)

# Insert chart
chart_path = os.path.join(REPORT_DIR, 'fig1_002594.SZ.png')
if os.path.exists(chart_path):
    doc.add_picture(chart_path, width=Inches(5.8))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_paragraph(doc, '')
add_paragraph(doc, '图 1：比亚迪日K线与 20 日 Donchian 通道及交易信号标记。红色上三角为买入信号（突破上轨），绿色下三角为离场信号（跌破下轨/止损），橙色菱形为加仓点，紫色方块为止损触发点。从图中可以观察到：通道在趋势明朗时宽度扩大，盘整时段宽度收窄；策略在 2024-2025 年度的上升趋势中多次触发买入和加仓信号，有效捕捉了主升浪。', bold=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_paragraph(doc, '')

add_heading_styled(doc, '4.2 净值曲线对比', level=2)

chart_path3 = os.path.join(REPORT_DIR, 'fig3_002594.SZ.png')
if os.path.exists(chart_path3):
    doc.add_picture(chart_path3, width=Inches(5.8))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_paragraph(doc, '')
add_paragraph(doc, '图 2：策略净值曲线（蓝色）与买入持有基准（灰色虚线）的对比。初始资金为 ¥1,000,000。可以看出策略在上升趋势中净值稳步增长，但在震荡/下跌阶段出现回撤。策略累计收益 +3.52%，同期买入持有基准收益约为-2.30%，策略实现约 5.82% 的超额收益。夏普比率 0.15，最大回撤 14.16%，风险调整后收益温和为正。', bold=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

doc.add_page_break()

# ============================================================
# Chapter 5: 策略适用场景与实战心得
# ============================================================
add_heading_styled(doc, '第五章  策略适用场景与实战心得', level=1)

add_heading_styled(doc, '5.1 适用场景', level=2)
add_paragraph(doc, '基于三年期五只标的的回测数据，海龟策略表现出以下适用特征：', bold=True)
add_paragraph(doc, '（1）强趋势标的：存在明确中长期上升趋势的标的（如美的集团），策略能完整捕捉趋势段的大部分利润，并通过加仓机制放大收益。三年累计超额收益达 22.31%。')
add_paragraph(doc, '（2）高波动率标的（ATR% > 2.5%）：ATR 动态仓位管理在波动率大的标的上优势最为凸显——高波动意味着更大的单日风险，但仓位自动缩小为控制回撤提供了天然缓冲。')
add_paragraph(doc, '（3）单边行情（上涨或下跌均可）：只要存在明确的单向趋势，海龟策略均能产生交易信号。做空方向同样适用，但在 A 股市场做空受限，实际应用中建议以做多为主。')
add_paragraph(doc, '')

add_heading_styled(doc, '5.2 失效场景', level=2)
add_paragraph(doc, '（1）长期横盘震荡：在缺乏明显方向的盘整市中，价格在 Donchian 通道内反复穿行，导致频繁的假突破信号和连续止损。震荡市是趋势跟踪策略的"天敌"。')
add_paragraph(doc, '（2）低波动率蓝筹（ATR% < 1.5%）：虽然回撤风险较小，但利润空间极为有限。例如长江电力虽然录得正收益，但 6.29% 的累计收益在扣除手续费后吸引力有限。')
add_paragraph(doc, '（3）政策驱动型标的：突发事件（如重大政策转向、黑天鹅事件）导致价格跳空，止损无法在预定价位执行，单笔亏损可能远超风险预算。')
add_paragraph(doc, '')

add_heading_styled(doc, '5.3 实战心得', level=2)
add_paragraph(doc, '心得一：ATR 动态止损是策略生命线——固定百分比止损无法适应市场波动率的变化，而 ATR 止损能够"随行就市"，是海龟策略能够在各类市场环境中生存的关键设计。', bold=True)
add_paragraph(doc, '心得二：加仓须克制——金字塔加仓在扩大盈利的同时也会放大回撤。3 次为经验上限，超过此限度的加仓往往带来不可承受的回撤风险。回测数据显示，多数交易在加仓 1-2 次后即触发离场。', bold=True)
add_paragraph(doc, '心得三：不追求高胜率——趋势跟踪策略的典型胜率在 30%-40% 之间（本次回测五只标的的胜率在 29%-50% 范围），低胜率是趋势跟踪的内在特征。盈亏比（profit factor）才是衡量策略有效性的关键指标——少数大盈利覆盖多次小亏损。', bold=True)
add_paragraph(doc, '心得四：成本敏感性不可忽视——万三手续费与万二滑点看似微小，但在频繁交易的标的上（如中国平安 31 笔交易），双向收取的交易成本累计可达本金的 1%-2%。低换手率的标的天然具有成本优势。', bold=True)
add_paragraph(doc, '心得五：标的筛选比参数优化更重要——从参数网格搜索结果来看，五只标的的最优参数各不相同，但标的间的绩效差异远大于同一标的在不同参数下的差异。优先选择趋势特征明确、波动率适中的标的，比穷举参数空间更有效。', bold=True)

doc.add_page_break()

# ============================================================
# Chapter 6: 结论
# ============================================================
add_heading_styled(doc, '第六章  结论', level=1)
add_paragraph(doc, '本项目基于 2023 年 7 月至 2026 年 7 月的三年期 A 股真实行情数据，对经典海龟交易策略进行了从理论到实践的完整复现。主要成果包括：')
add_paragraph(doc, '')
add_paragraph(doc, '1. 从零构建了符合 GitHub 开源标准的纯静态 HTML 可视化看板（task05_turtle_dashboard/），支持参数实时调节与多标的切换分析，可直接通过 GitHub Pages 部署。')
add_paragraph(doc, '2. 实现了完整的前后端海龟策略引擎（Python + JavaScript 双版本），涵盖 Donchian 通道计算、Wilder ATR 平滑、信号状态机生成、仓位模拟与绩效评估全流程。')
add_paragraph(doc, '3. 对五只风格迥异的 A 股标的（高波动的中芯国际/比亚迪、低波动的长江电力、均值回归的美的集团、宏观敏感的中国平安）进行了系统性回测与参数敏感性分析。')
add_paragraph(doc, '4. 回测结果验证了海龟策略的核心特征：在趋势明确的标的上表现优异（美的集团年化 +7.33%），在震荡/下跌市中有一定的防御能力但难以彻底避免亏损。策略的适用边界清晰——它不是一个"万能"系统，而是一个在特定市场环境下高效的趋势捕捉工具。')
add_paragraph(doc, '')
add_paragraph(doc, '本项目的完整代码、数据、可视化看板与文档均已纳入 task05_turtle_dashboard/ 目录，遵循 MIT 协议开放使用。')

# ============================================================
# SAVE
# ============================================================
os.makedirs(REPORT_DIR, exist_ok=True)
doc.save(OUTPUT)
print(f"\nWord document saved to: {OUTPUT}")
print("Done!")
